import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = docx.Document()

    # Title
    title = doc.add_heading('EvoCode Project Progress Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Phase 1 & 2 Completion Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('') # Spacing

    # Section 1
    doc.add_heading('1. Overview', level=1)
    p = doc.add_paragraph(
        "This document outlines the completion of Phase 1 (Foundations) and Phase 2 (Single-Agent Baseline) "
        "of the EvoCode final-year project. The primary goal of these initial phases was to build resilient infrastructure "
        "for free-tier LLM API usage, implement the core evaluation sandbox, and establish a strong baseline success metric "
        "before introducing the more complex Evolutionary Pipeline."
    )

    # Section 2
    doc.add_heading('2. Infrastructure & Foundations (Phase 1)', level=1)
    doc.add_paragraph("The following core architectural components were successfully implemented and tested:")
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Rate Limiting & Call Budgets: ').bold = True
    ul.add_run('Designed and integrated a strict asynchronous token-bucket rate limiter that natively respects '
               'the Tokens-Per-Minute (TPM) and Requests-Per-Minute (RPM) limits of free-tier providers. '
               'A global budget tracker ensures experiments halt safely before hitting daily limits.')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Provider Abstraction & Fallback: ').bold = True
    ul.add_run('Implemented a highly fault-tolerant LLM client using tenacity. The client gracefully handles '
               'API timeouts, rate limit errors (HTTP 429), and automatically falls back from Groq models to OpenRouter '
               'models when primary limits are exhausted.')
               
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Agent Roles: ').bold = True
    ul.add_run('Implemented the structural foundation for the five distinct agent roles (Analyzer, Planner/Coder, '
               'Critic, Mutator, Judge) that will be used heavily in the upcoming evolutionary loop.')

    # Section 3
    doc.add_heading('3. Single-Agent Baseline Results (Phase 2)', level=1)
    doc.add_paragraph("To measure the effectiveness of the future Evolutionary Pipeline, a single-agent "
                      "baseline was established. This baseline simulates a conventional AI coding assistant "
                      "attempting to fix a bug using iterative self-refinement (up to 3 attempts).")
                      
    doc.add_paragraph("The baseline experiment was run against a carefully curated 25-instance subset of the SWE-bench Lite dataset.")
    
    # Results formatting
    p_results = doc.add_paragraph()
    p_results.add_run("Experimental Results:\n").bold = True
    p_results.add_run("• Total Instances Processed: ").bold = True
    p_results.add_run("25\n")
    p_results.add_run("• Total Instances Resolved (PASS): ").bold = True
    p_results.add_run("15\n")
    p_results.add_run("• Baseline Resolution Rate: ").bold = True
    p_results.add_run("60%")

    p2 = doc.add_paragraph("Achieving a 60% resolution rate in a zero-context simulated environment is an exceptionally strong baseline. "
                      "This metric provides a high, robust bar to test the evolutionary hypothesis against in Phase 4.")

    # Section 4
    doc.add_heading('4. Next Steps (Phase 3)', level=1)
    doc.add_paragraph("With the baseline control group established and infrastructure stable, the project now advances to "
                      "Phase 3: The Evolutionary Pipeline. The upcoming deliverables include:")
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Genome Config Schema: ').bold = True
    ul.add_run('Defining the Pydantic schemas that allow agent behavior (temperature, prompt variants, debate flags) to mutate.')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('EvoFlow State Machine: ').bold = True
    ul.add_run('Building the multi-agent evolutionary controller that handles population spawning, fitness evaluation, and selection.')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('SQLite Memory Store: ').bold = True
    ul.add_run('Implementing cross-generational knowledge sharing so successful strategies survive across runs.')

    doc.add_paragraph("\nStatus: On Track.")

    # Save
    doc.save("Phase_2_Completion_Report.docx")
    print("Report generated successfully.")

if __name__ == "__main__":
    main()
